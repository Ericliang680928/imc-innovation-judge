"""
評審報告產生器:把 Claude 回傳的 Markdown 轉成 Word 與 PDF 下載檔。
"""

import io
import re
from datetime import datetime
from pathlib import Path


# ─────────────────────────────────────────────
# Word 報告
# ─────────────────────────────────────────────
def markdown_to_docx(markdown_text: str, meta: dict) -> bytes:
    """
    把 Markdown 評審報告轉成 .docx,回傳 bytes。
    meta: {'category': str, 'stage': str, 'case_title': str, 'file_name': str}
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 設定預設字型(讓中文用「微軟正黑體」)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)
    try:
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    except Exception:
        pass

    # 標題頁資訊
    title = doc.add_heading("IMC 第 21 屆創新獎 評審報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"參賽類組:{meta.get('category', '')}　|　評審階段:{meta.get('stage', '')}\n"
        f"檔案來源:{meta.get('file_name', '')}\n"
        f"產生時間:{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 解析 Markdown 並轉為 docx 元素
    lines = markdown_text.split("\n")
    in_table = False
    table_buffer = []

    for raw in lines:
        line = raw.rstrip()

        # 表格偵測(以 | 開頭結尾)
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            in_table = True
            continue
        else:
            if in_table and table_buffer:
                _add_md_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not line.strip():
            doc.add_paragraph()
            continue

        # 程式碼區塊邊界(忽略 ``` 行)
        if line.strip().startswith("```"):
            continue

        # 分隔線
        if re.match(r"^[-*_]{3,}$", line.strip()):
            p = doc.add_paragraph()
            p.add_run("─" * 50)
            continue

        # 標題
        h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(_strip_md(h.group(2)), level=level)
            continue

        # 編號清單
        m_ol = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m_ol.group(2))
            continue

        # 項目清單
        m_ul = re.match(r"^[-*]\s+(.+)$", line)
        if m_ul:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m_ul.group(1))
            continue

        # 縮排項目(  - 或 - )
        m_indent = re.match(r"^(\s{2,})[-*]\s+(.+)$", line)
        if m_indent:
            p = doc.add_paragraph(style="List Bullet 2")
            _add_inline_runs(p, m_indent.group(2))
            continue

        # 一般段落
        p = doc.add_paragraph()
        _add_inline_runs(p, line)

    # 收尾表格
    if in_table and table_buffer:
        _add_md_table(doc, table_buffer)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_md_table(doc, table_buffer):
    """把 markdown 表格行轉成 docx 表格。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    rows = []
    for line in table_buffer:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # 略過分隔列(--- 列)
        if all(re.match(r"^:?-{3,}:?$", c) for c in cells if c):
            continue
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    # 補齊欄位數
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(_strip_md(cell_text))
            run.font.size = Pt(10)
            try:
                run.font.name = "Microsoft JhengHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            except Exception:
                pass
            if i == 0:
                run.bold = True


def _add_inline_runs(paragraph, text):
    """處理粗體 ** ** 與斜體,並設定中文字型。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # 切分粗體
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        clean = _strip_md(clean, keep_bold=False)
        run = paragraph.add_run(clean)
        run.font.size = Pt(11)
        if is_bold:
            run.bold = True
        try:
            run.font.name = "Microsoft JhengHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        except Exception:
            pass


def _strip_md(text, keep_bold=True):
    """移除其他 markdown 標記。"""
    if not keep_bold:
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


# ─────────────────────────────────────────────
# PDF 報告(用 ReportLab + 系統中文字型)
# ─────────────────────────────────────────────
def markdown_to_pdf(markdown_text: str, meta: dict) -> bytes:
    """把 Markdown 評審報告轉成 PDF,回傳 bytes。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 註冊中文字型
    font_name = _register_cjk_font()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title="IMC 第21屆創新獎評審報告",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN", parent=styles["Title"], fontName=font_name, fontSize=20,
        alignment=TA_CENTER, spaceAfter=10, textColor=colors.HexColor("#1a4d7a")
    )
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name,
                        fontSize=15, textColor=colors.HexColor("#1a4d7a"),
                        spaceBefore=12, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=font_name,
                        fontSize=13, textColor=colors.HexColor("#2a6ca8"),
                        spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=font_name,
                        fontSize=11.5, textColor=colors.HexColor("#444444"),
                        spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=font_name,
                          fontSize=10.5, leading=16, alignment=TA_LEFT, spaceAfter=4)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=18, bulletIndent=4)
    meta_style = ParagraphStyle("Meta", parent=body, fontSize=9.5,
                                textColor=colors.HexColor("#666666"),
                                alignment=TA_CENTER, spaceAfter=12)

    story = []
    story.append(Paragraph("IMC 第 21 屆創新獎 評審報告", title_style))
    story.append(Paragraph(
        f"參賽類組:{meta.get('category', '')}　|　評審階段:{meta.get('stage', '')}<br/>"
        f"檔案來源:{_escape_html(meta.get('file_name', ''))}<br/>"
        f"產生時間:{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        meta_style
    ))
    story.append(Spacer(1, 0.4 * cm))

    # 解析 Markdown
    lines = markdown_text.split("\n")
    table_buffer = []
    in_table = False

    for raw in lines:
        line = raw.rstrip()

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            in_table = True
            continue
        else:
            if in_table and table_buffer:
                story.append(_make_pdf_table(table_buffer, font_name))
                story.append(Spacer(1, 0.2 * cm))
                table_buffer = []
                in_table = False

        if not line.strip():
            story.append(Spacer(1, 0.15 * cm))
            continue

        if line.strip().startswith("```"):
            continue

        if re.match(r"^[-*_]{3,}$", line.strip()):
            story.append(Spacer(1, 0.1 * cm))
            continue

        h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h:
            level = len(h.group(1))
            text = _md_inline_to_html(h.group(2))
            style_map = {1: h1, 2: h1, 3: h2, 4: h3}
            story.append(Paragraph(text, style_map.get(level, h3)))
            continue

        m_ol = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m_ol:
            text = f"{m_ol.group(1)}. {_md_inline_to_html(m_ol.group(2))}"
            story.append(Paragraph(text, bullet))
            continue

        m_ul = re.match(r"^[-*]\s+(.+)$", line)
        if m_ul:
            text = f"• {_md_inline_to_html(m_ul.group(1))}"
            story.append(Paragraph(text, bullet))
            continue

        m_indent = re.match(r"^\s{2,}[-*]\s+(.+)$", line)
        if m_indent:
            text = f"&nbsp;&nbsp;&nbsp;◦ {_md_inline_to_html(m_indent.group(1))}"
            story.append(Paragraph(text, bullet))
            continue

        story.append(Paragraph(_md_inline_to_html(line), body))

    if in_table and table_buffer:
        story.append(_make_pdf_table(table_buffer, font_name))

    doc.build(story)
    return buf.getvalue()


def _make_pdf_table(table_buffer, font_name):
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    rows = []
    for line in table_buffer:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^:?-{3,}:?$", c) for c in cells if c):
            continue
        rows.append(cells)

    if not rows:
        return Paragraph("", ParagraphStyle("empty"))

    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    cell_style = ParagraphStyle(
        "TableCell", fontName=font_name, fontSize=9.5,
        leading=13, alignment=TA_LEFT
    )
    head_style = ParagraphStyle(
        "TableHead", fontName=font_name, fontSize=10,
        leading=13, alignment=TA_LEFT, textColor=colors.white
    )

    data = []
    for i, row in enumerate(rows):
        if i == 0:
            data.append([Paragraph(_md_inline_to_html(c), head_style) for c in row])
        else:
            data.append([Paragraph(_md_inline_to_html(c), cell_style) for c in row])

    # 自動分配欄寬
    from reportlab.lib.units import cm
    total_w = 17 * cm
    if n_cols == 4:
        col_widths = [total_w * 0.30, total_w * 0.12, total_w * 0.12, total_w * 0.46]
    elif n_cols == 3:
        col_widths = [total_w * 0.35, total_w * 0.15, total_w * 0.50]
    else:
        col_widths = [total_w / n_cols] * n_cols

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a6ca8")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#888888")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f7fb")]),
    ]))
    return t


def _md_inline_to_html(text):
    """把 markdown inline 標記轉成 ReportLab 的 mini-HTML。"""
    text = _escape_html(text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r'<font face="Courier">\1</font>', text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def _escape_html(text):
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))


def _register_cjk_font():
    """註冊中文字型,優先順序:微軟正黑體 → 標楷體 → 思源黑體 → 內建 CID 字型。"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    candidates = [
        ("MSJhengHei", r"C:\Windows\Fonts\msjh.ttc"),
        ("MSJhengHei", r"C:\Windows\Fonts\msjh.ttf"),
        ("MingLiU", r"C:\Windows\Fonts\mingliu.ttc"),
        ("DFKai-SB", r"C:\Windows\Fonts\kaiu.ttf"),
        ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for name, path in candidates:
        if Path(path).exists():
            try:
                # ttc 需指定 subfontIndex
                if path.lower().endswith(".ttc"):
                    pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue

    # fallback: ReportLab 內建 CJK
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"
